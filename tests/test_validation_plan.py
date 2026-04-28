"""バリデーションプラン (Word) 生成テスト。"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.generators.validation_plan import write_validation_plan
from src.parser.edc_parser import load_study

SAMPLE_JSON = Path(
    r"c:/Users/matsu/Box/Datacenter/ISR/Ptosh/検証/JSON/Bev-FOLFOX-SBC_250929_1501.json"
)


@pytest.fixture(scope="module")
def vp_path(tmp_path_factory) -> Path:
    if not SAMPLE_JSON.exists():
        pytest.skip("sample JSON not found")
    study = load_study(SAMPLE_JSON)
    out = tmp_path_factory.mktemp("vp") / "validation_plan.docx"
    return write_validation_plan(study, out)


def test_headings_present(vp_path: Path):
    doc = Document(vp_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    expected = [
        "1. 文書管理", "2. はじめに", "3. 試験概要", "4. EDC システム概要",
        "5. バリデーション対象範囲", "6. バリデーション方法",
        "7. テスト方針", "8. 不適合管理", "9. 役割と責任", "10. 添付資料",
    ]
    for h in expected:
        assert h in headings, f"missing heading: {h}"
    # 9.3 (システム担当) は削除済
    assert not any("9.3" in t for t in headings)


def test_study_meta_in_kv(vp_path: Path):
    """試験 ID と試験名称が表に含まれること。"""
    doc = Document(vp_path)
    text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Bev-FOLFOX-SBC" in text
    assert "小腸癌" in text


def test_validator_counts(vp_path: Path):
    doc = Document(vp_path)
    text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    # 6章のロジカルチェック数 4,065
    assert "4,065" in text
    # 5章のフォーム数 162
    assert "162" in text
