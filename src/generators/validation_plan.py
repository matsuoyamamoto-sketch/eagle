"""バリデーションプラン (Word) 生成モジュール。

確定目次に基づき、python-docx でプログラム的に文書を組み立てる。
(汎用版テンプレートは外部ファイル化せず、本モジュール内に集約。)
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from ..parser.models import Study

BASE_FONT = "Meiryo UI"


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BASE_FONT
    style.font.size = Pt(10.5)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def _add_h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = BASE_FONT
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def _add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = BASE_FONT


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """2列の Key/Value テーブル。"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        kc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        kc.text = k
        vc.text = str(v)
        kc.width = Cm(5)
        vc.width = Cm(11)
        for para in kc.paragraphs:
            for run in para.runs:
                run.font.name = BASE_FONT
                run.font.bold = True
        for para in vc.paragraphs:
            for run in para.runs:
                run.font.name = BASE_FONT


def _add_count_table(doc: Document, rows: list[tuple[str, int]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "項目"
    hdr[1].text = "件数"
    for i, (k, v) in enumerate(rows, start=1):
        c = table.rows[i].cells
        c[0].text = k
        c[1].text = f"{v:,}"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = BASE_FONT


# ---------------- 各章 ----------------
def _ch1_document_management(doc: Document) -> None:
    _add_h(doc, "1. 文書管理", level=1)
    _add_h(doc, "1.1 改訂履歴", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["版", "日付", "改訂者", "改訂内容"]):
        hdr[i].text = h
    row = table.rows[1].cells
    row[0].text = "1.0"
    row[1].text = date.today().isoformat()
    row[2].text = ""
    row[3].text = "初版作成 (EAGLE 自動生成)"

    _add_h(doc, "1.2 承認欄", level=2)
    sign = doc.add_table(rows=2, cols=3)
    sign.style = "Light Grid Accent 1"
    for i, h in enumerate(["作成", "確認", "承認"]):
        sign.rows[0].cells[i].text = h
        sign.rows[1].cells[i].text = ""


def _ch2_introduction(doc: Document, study: Study) -> None:
    _add_h(doc, "2. はじめに", level=1)
    _add_h(doc, "2.1 目的", level=2)
    _add_para(
        doc,
        f"本書は、試験「{study.name}」における EDC システムの設定が、"
        "プロトコルおよび関連手順書に整合し、データ品質要件を満たしていることを"
        "確認するためのバリデーション計画を定めるものである。",
    )
    _add_h(doc, "2.2 適用範囲", level=2)
    _add_para(
        doc,
        "本計画は、当該試験用に構築された EDC システムの全フォーム、入力項目、"
        "コードリスト、エディットチェックを対象とする。",
    )
    _add_h(doc, "2.3 関連文書", level=2)
    _add_para(doc, "・プロトコル")
    _add_para(doc, "・CRF 設計書")
    _add_para(doc, "・データマネジメント計画書 (DMP)")
    _add_para(doc, "・本システムの EDC 仕様書 (別添)")


def _ch3_study_overview(doc: Document, study: Study) -> None:
    _add_h(doc, "3. 試験概要", level=1)
    _add_kv_table(
        doc,
        [
            ("試験 ID", study.name),
            ("試験名称", study.proper_name),
            ("実施機関", study.organization_name or "-"),
            ("疾患カテゴリ", study.disease_category or "-"),
            ("SDTM Version", study.sdtm_version or "-"),
            ("SDTM Terminology", study.sdtm_terminology_version or "-"),
            ("CTCAE Version", study.ctcae_version or "-"),
        ],
    )


def _ch4_edc_system(doc: Document) -> None:
    _add_h(doc, "4. EDC システム概要", level=1)
    _add_h(doc, "4.1 構成", level=2)
    _add_para(doc, "Web ブラウザを介してアクセスする EDC システム。サーバ・データベース・ "
                   "認証基盤の構成は別途インフラ仕様書を参照する。")
    _add_h(doc, "4.2 アクセス権限", level=2)
    _add_para(doc, "ロール (CRC / Investigator / Monitor / DM / Auditor 等) ごとに権限を分離する。")
    _add_h(doc, "4.3 監査証跡", level=2)
    _add_para(doc, "全データ更新は ALCOA+ 原則に従い、変更前後値・実施者・日時を保持する。")


def _ch5_scope(doc: Document, study: Study) -> None:
    _add_h(doc, "5. バリデーション対象範囲", level=1)
    used_codelists = len({fi.option_name for s in study.sheets for fi in s.field_items if fi.option_name})
    _add_count_table(
        doc,
        [
            ("フォーム数", len(study.sheets)),
            ("総入力項目数", study.total_field_items()),
            ("コードリスト数 (使用中)", used_codelists),
            ("割付群数", len(study.sheet_groups)),
        ],
    )
    if study.sheet_groups:
        _add_para(doc, "割付群: " + " / ".join(g.name for g in study.sheet_groups))


def _ch6_strategy(doc: Document, study: Study) -> None:
    _add_h(doc, "6. バリデーション方法", level=1)
    _add_count_table(
        doc,
        [
            ("必須項目チェック (presence)", study.count_validator("presence")),
            ("日付範囲チェック (date)", study.count_validator("date")),
            ("数値範囲チェック (numericality)", study.count_validator("numericality")),
            ("ロジカルチェック (formula)", study.count_validator("formula")),
            ("シート間参照項目 (Reference)", study.count_reference_items()),
        ],
    )
    _add_para(doc, "上記の各チェックについて、別添「エディットチェック確認書」に")
    _add_para(doc, "条件式およびエラーメッセージの全件を記載する。")
    _add_h(doc, "6.5 コードリスト整合性", level=2)
    _add_para(doc, "各 radio_button 項目に紐づくコードリストの値が CDISC SDTM 標準に準拠していることを確認する。")
    _add_h(doc, "6.6 シート間参照整合性", level=2)
    _add_para(doc, "FieldItem::Reference により他フォームの値を引用する項目について、参照元と一致することを確認する。")


def _ch7_test_strategy(doc: Document) -> None:
    _add_h(doc, "7. テスト方針", level=1)
    _add_h(doc, "7.1 単体テスト", level=2)
    _add_para(doc, "境界値分析および同値分割により、各入力項目の正常系・異常系入力を網羅する。")
    _add_h(doc, "7.2 結合テスト", level=2)
    _add_para(doc, "シート間参照、割付処理、合計値計算等のフォーム横断ロジックを検証する。")
    _add_h(doc, "7.3 ユーザー受入テスト (UAT)", level=2)
    _add_para(doc, "実運用に近いシナリオで、CRC・モニター・DM 担当者が一連の操作を実施し承認する。")


def _ch8_nonconformance(doc: Document) -> None:
    _add_h(doc, "8. 不適合管理", level=1)
    _add_h(doc, "8.1 検出時の対応フロー", level=2)
    _add_para(doc, "テスト中に不適合が検出された場合、Issue Tracker に登録し、原因解析・修正・再テストを実施する。")
    _add_h(doc, "8.2 再テスト基準", level=2)
    _add_para(doc, "修正対象のチェック項目、および影響を受ける関連項目について再テストを実施する。")


def _ch9_roles(doc: Document) -> None:
    _add_h(doc, "9. 役割と責任", level=1)
    _add_h(doc, "9.1 DM 担当", level=2)
    _add_para(doc, "本計画の策定、テスト実施、不適合管理、最終承認を行う。")
    _add_h(doc, "9.2 統計担当", level=2)
    _add_para(doc, "解析データセット要件との整合性を確認し、必要なチェック項目を提案する。")


def _ch10_attachments(doc: Document) -> None:
    _add_h(doc, "10. 添付資料", level=1)
    _add_para(doc, "・EDC 仕様書 (別添 Excel)")
    _add_para(doc, "・エディットチェック確認書 (別添 Excel)")
    _add_para(doc, "・テストシナリオ (別添 Excel)")
    _add_para(doc, "・マニュアルチェックリスト (別添 Excel)")


# ---------------- エントリ ----------------
def build_validation_plan(study: Study) -> Document:
    doc = Document()
    _set_default_style(doc)

    # 表紙
    for _ in range(4):
        doc.add_paragraph("")
    _add_title(doc, "バリデーションプラン")
    doc.add_paragraph("")
    _add_subtitle(doc, study.proper_name)
    doc.add_paragraph("")
    doc.add_paragraph("")
    _add_kv_table(doc, [("試験 ID", study.name), ("発行日", date.today().isoformat())])
    doc.add_page_break()

    _ch1_document_management(doc)
    _ch2_introduction(doc, study)
    _ch3_study_overview(doc, study)
    _ch4_edc_system(doc)
    _ch5_scope(doc, study)
    _ch6_strategy(doc, study)
    _ch7_test_strategy(doc)
    _ch8_nonconformance(doc)
    _ch9_roles(doc)
    _ch10_attachments(doc)
    return doc


def write_validation_plan(study: Study, output_path: str | Path) -> Path:
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_validation_plan(study)
    doc.save(out)
    return out
